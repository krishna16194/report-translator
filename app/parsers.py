"""Extract a simple, format-agnostic content model from supported input files.

The content model is a list of "blocks". Each block is a dict with a "type":
  - {"type": "heading", "text": str, "level": int}
  - {"type": "paragraph", "text": str}
  - {"type": "table", "rows": list[list[str]]}

This intermediate representation lets the translator and the document
generators stay independent of the input format.
"""
from __future__ import annotations

import os
from typing import Any

from docx import Document as DocxDocument
from openpyxl import load_workbook


class UnsupportedFileError(ValueError):
    """Raised when an uploaded file has an extension we cannot parse."""


def parse(file_path: str) -> dict[str, Any]:
    """Parse a file into {"title": str, "blocks": [...]}.

    Dispatches on the file extension.
    """
    ext = os.path.splitext(file_path)[1].lower()
    base_title = os.path.splitext(os.path.basename(file_path))[0]

    if ext == ".txt":
        blocks = _parse_txt(file_path)
    elif ext == ".docx":
        blocks = _parse_docx(file_path)
    elif ext in (".xlsx", ".xlsm"):
        blocks = _parse_xlsx(file_path)
    else:
        raise UnsupportedFileError(
            f"Unsupported file type '{ext}'. Supported: .txt, .docx, .xlsx"
        )

    return {"title": base_title, "blocks": blocks}


def _parse_txt(file_path: str) -> list[dict[str, Any]]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as fh:
        text = fh.read()

    blocks: list[dict[str, Any]] = []
    # Split on blank lines into paragraphs; keep non-empty chunks.
    for chunk in text.split("\n\n"):
        chunk = chunk.strip()
        if chunk:
            blocks.append({"type": "paragraph", "text": chunk})
    return blocks


def _parse_docx(file_path: str) -> list[dict[str, Any]]:
    doc = DocxDocument(file_path)
    blocks: list[dict[str, Any]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            # Style names look like "Heading 1", "Heading 2", ...
            level = 1
            digits = "".join(ch for ch in style if ch.isdigit())
            if digits:
                level = int(digits)
            blocks.append({"type": "heading", "text": text, "level": level})
        elif style == "title":
            blocks.append({"type": "heading", "text": text, "level": 1})
        else:
            blocks.append({"type": "paragraph", "text": text})

    for table in doc.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            blocks.append({"type": "table", "rows": rows})

    return blocks


def _parse_xlsx(file_path: str) -> list[dict[str, Any]]:
    wb = load_workbook(file_path, data_only=True, read_only=True)
    blocks: list[dict[str, Any]] = []

    for sheet in wb.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            # Skip fully empty rows.
            if any(c.strip() for c in cells):
                rows.append(cells)
        if rows:
            blocks.append({"type": "heading", "text": sheet.title, "level": 2})
            blocks.append({"type": "table", "rows": rows})

    wb.close()
    return blocks
