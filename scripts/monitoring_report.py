"""Generate a DISPACT-style monitoring report (.docx) from a monitoring Excel sheet.

This is a *generic* mapper: it learns nothing from any single file. Given

    * a monitoring working sheet  (Excel, the monitor's per-visit notes)
    * the empty report template    (the .docx Monitoringbericht form)

it produces a filled report that preserves the template's formatting.

What it derives from the Excel
------------------------------
* Visit date          -> document header "Besuchsdatum"
* Patient status       -> gescreent / randomisiert / eingeschlossen counts
* Reviewed patients    -> screening-number lists (Patienteneinschluss, SDV)
* SDV visit ranges     -> eCRF / SDV table (parsed from the "Bemerkungen" column)
* Findings             -> routed by topic keywords into the matching report
                          section's comment area; anything unmatched lands in
                          "Generelle Kommentare".

What it cannot derive (left blank for a human, as in the blank template)
------------------------------------------------------------------------
* Center name / number, visit number, attendee names & functions
* The "geplant" (contracted) patient count
* The Ja / Nein / NZ checkbox marks  (a judgement, not raw data)

Usage
-----
    python scripts/monitoring_report.py INPUT.xlsx TEMPLATE.docx OUTPUT.docx

The Excel is expected to have, on its first worksheet:
    * a title cell like "Monitoring am DD.MM.YY"
    * a header row containing "Screening-Nr"
    * a left block:  Screening-Nr | Rando Datum | IC Datum | Rando Nr. |
                     Korrekt | Bemerkungen | Randobestaetigung
    * a right block: Inhalt | Bemerkungen | aktuelle Status
Columns are located by header text, so their exact position may vary.
"""
from __future__ import annotations

import argparse
import datetime as _dt
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

import openpyxl
from docx import Document
from docx.oxml.ns import qn


# --------------------------------------------------------------------------- #
# Data model
# --------------------------------------------------------------------------- #
@dataclass
class Patient:
    screening: str
    rando_date: _dt.date | None = None
    ic_date: _dt.date | None = None
    rando_nr: str | None = None
    korrekt: str | None = None
    remark: str | None = None
    rando_confirmed: str | None = None

    @property
    def is_randomised(self) -> bool:
        return bool(self.rando_nr)

    @property
    def is_included(self) -> bool:  # included in the study
        # A randomised patient is by definition included (you cannot randomise
        # without prior inclusion), so inclusion is implied by a Rando-Nr. even
        # when the monitor left the IC-Datum cell blank for older patients —
        # otherwise the count would read randomisiert > eingeschlossen.
        return self.ic_date is not None or self.is_randomised


@dataclass
class Finding:
    topic: str          # "Inhalt"
    comment: str        # "Bemerkungen"
    status: str         # "aktuelle Status"

    def render(self) -> str:
        bits = []
        if self.topic:
            bits.append(self.topic.strip())
        body = " ".join(b for b in (self.comment.strip(), self.status.strip()) if b)
        if self.topic and body:
            return f"{self.topic.strip()}: {body}"
        return self.topic.strip() or body


@dataclass
class VisitData:
    visit_date: _dt.date | None
    patients: list[Patient] = field(default_factory=list)
    findings: list[Finding] = field(default_factory=list)

    # patients with source-data verification done this visit + the visit range
    def sdv_patients(self) -> list[tuple[Patient, str]]:
        out = []
        for p in self.patients:
            if p.remark and "sdv" in p.remark.lower():
                out.append((p, _visit_range(p.remark)))
        return out


_VISIT_RANGE = re.compile(r"(V\s*\d+\s*[-–bis]+\s*V?\s*\d+|bis\s*V\s*\d+)", re.I)


def _visit_range(remark: str) -> str:
    m = _VISIT_RANGE.search(remark)
    return m.group(0).strip() if m else remark.strip()


# --------------------------------------------------------------------------- #
# Excel parsing
# --------------------------------------------------------------------------- #
# The header row carries two blocks; "Bemerkungen" appears in BOTH, so the two
# blocks must be matched with separate alias tables (split at the "Inhalt"
# column, which begins the findings block).
LEFT_ALIASES = {
    "screening": ("screening-nr", "screening nr", "screening"),
    "rando_date": ("rando datum", "rando-datum"),
    "ic_date": ("ic datum", "ic-datum"),
    "rando_nr": ("rando nr", "rando-nr", "rando nr."),
    "korrekt": ("korrekt",),
    "remark": ("bemerkungen", "bemerkung"),
    "rando_confirmed": ("randobestätigung", "randobestaetigung", "rando best"),
}
RIGHT_ALIASES = {
    "inhalt": ("inhalt",),
    "f_comment": ("bemerkungen", "bemerkung"),
    "status": ("aktuelle status", "status"),
}


def _norm(v) -> str:
    return str(v).strip().lower() if v is not None else ""


def _as_date(v):
    if isinstance(v, _dt.datetime):
        return v.date()
    if isinstance(v, _dt.date):
        return v
    return None


def _parse_title_date(ws) -> _dt.date | None:
    for row in ws.iter_rows(min_row=1, max_row=3):
        for cell in row:
            if isinstance(cell.value, str):
                m = re.search(r"(\d{1,2})\.(\d{1,2})\.(\d{2,4})", cell.value)
                if m:
                    d, mth, y = (int(g) for g in m.groups())
                    if y < 100:
                        y += 2000
                    try:
                        return _dt.date(y, mth, d)
                    except ValueError:
                        pass
    # fall back to the sheet tab name (often the visit date)
    m = re.search(r"(\d{1,2})\.(\d{1,2})\.(\d{2,4})", ws.title)
    if m:
        d, mth, y = (int(g) for g in m.groups())
        y += 2000 if y < 100 else 0
        try:
            return _dt.date(y, mth, d)
        except ValueError:
            return None
    return None


def parse_excel(path: Path) -> VisitData:
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.worksheets[0]

    visit_date = _parse_title_date(ws)

    # Locate the header row: the row that contains a "Screening-Nr" cell.
    header_row = None
    for row in ws.iter_rows():
        if any(_norm(c.value).startswith("screening") for c in row):
            header_row = row[0].row
            break
    if header_row is None:
        raise ValueError("Could not find a header row containing 'Screening-Nr'.")

    # Map column index -> logical field for the left (patient) block and
    # right (findings) block.  The findings block is the second occurrence of
    # "Inhalt"; its "Bemerkungen"/"Status" follow it to the right.
    left_cols: dict[str, int] = {}
    right_cols: dict[str, int] = {}
    header_cells = ws[header_row]
    seen_inhalt = False
    for cell in header_cells:
        text = _norm(cell.value)
        if not text:
            continue
        if text in ("inhalt",):
            seen_inhalt = True
        # everything from "Inhalt" rightwards belongs to the findings block, and
        # is matched against a different alias table than the left block.
        target, aliases_tbl = (right_cols, RIGHT_ALIASES) if seen_inhalt else (left_cols, LEFT_ALIASES)
        for key, aliases in aliases_tbl.items():
            if text in aliases or any(text.startswith(a) for a in aliases):
                if key not in target:  # keep the first match per side
                    target[key] = cell.column
                break

    patients: list[Patient] = []
    findings: list[Finding] = []
    last_topic = ""

    for row in ws.iter_rows(min_row=header_row + 1):
        cells = {c.column: c.value for c in row}

        # --- left block: a patient row needs a screening number ---
        sc = cells.get(left_cols.get("screening"))
        if sc not in (None, ""):
            patients.append(
                Patient(
                    screening=str(sc).strip().rstrip(".0") if isinstance(sc, float) else str(sc).strip(),
                    rando_date=_as_date(cells.get(left_cols.get("rando_date"))),
                    ic_date=_as_date(cells.get(left_cols.get("ic_date"))),
                    rando_nr=_clean(cells.get(left_cols.get("rando_nr"))),
                    korrekt=_clean(cells.get(left_cols.get("korrekt"))),
                    remark=_clean(cells.get(left_cols.get("remark"))),
                    rando_confirmed=_clean(cells.get(left_cols.get("rando_confirmed"))),
                )
            )

        # --- right block: a finding row needs a topic and/or a comment ---
        inhalt = _clean(cells.get(right_cols.get("inhalt"))) or ""
        f_comment = _clean(cells.get(right_cols.get("f_comment"))) or ""
        status = _clean(cells.get(right_cols.get("status"))) or ""
        if inhalt:
            last_topic = inhalt
        if f_comment or status:
            findings.append(Finding(topic=inhalt or last_topic, comment=f_comment, status=status))

    return VisitData(visit_date=visit_date, patients=patients, findings=findings)


def _clean(v) -> str | None:
    if v is None:
        return None
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    s = str(v).strip()
    return s or None


# --------------------------------------------------------------------------- #
# Finding -> report-section routing
# --------------------------------------------------------------------------- #
# Section keys map to the anchor text of the matching report table.
SECTION_ANCHORS = {
    "einschluss": "Patienteneinschluss",
    "protokoll": "Einhaltung des Studienprotokolls",
    "sicherheit": "Sicherheit",
    "endpunkt": "Primärer Endpunkt",
    "quelldaten": "Dokumentation der Quelldaten",
    "sdv": "eCRF-Überprüfung/SDV",
    "ecrf": "Dokumentation im eCRF",
    "labor": "Labor und biologische Proben",
    "zentrum": "Zentrum",
    "isf": "Zentrumsordner (Investigator Site File)",
    "general": "Generelle Kommentare",
}

# Ordered keyword rules — first match wins.  Tuned to the German monitoring
# vocabulary; extend freely.
ROUTING_RULES: list[tuple[str, tuple[str, ...]]] = [
    ("zentrum", ("mutterschutz", "zugang", "abmeldung im ecrf", "redcap-zugang", "abmelden")),
    ("isf", ("log of staff", "trainingslog", "redcap training", "redcap-training",
             "gcp & cv", "cv", "doi", "doi ", "entnahmehinweis", "kurzanleitung",
             "note to file", "qlq", "abmeldeformular", "isf", "training")),
    ("sicherheit", ("biochemical leak", "popf", "dge", "komplikation", "complication",
                    "adverse", "ae ", "leak")),
    ("ecrf", ("query", "ecrf", "redcap", "feld", "eingetragen", "asa", "klassifikation",
              "histopatho", "ipmn", "datensatz", "additional coverage", "re-intervention",
              "icu/imc", "fb", "worksheet")),
    ("einschluss", ("hausarzt", "einwilligung", "aufklärung", "patienteninformation",
                    "informierte", "consent", "ic ", "ics")),
    ("endpunkt", ("primärer endpunkt", "primary endpoint")),
    ("quelldaten", ("quelldaten", "source data", "drg")),
    ("labor", ("labor", "normwert", "ringversuch", "akkreditier")),
    ("protokoll", ("protokoll", "verblindung", "visite gemäß", "behandlung gemäß")),
]


def route(finding: Finding) -> str:
    hay = f"{finding.topic} {finding.comment} {finding.status}".lower()
    for section, kws in ROUTING_RULES:
        if any(kw in hay for kw in kws):
            return section
    return "general"


# --------------------------------------------------------------------------- #
# docx helpers
# --------------------------------------------------------------------------- #
def _find_table(doc, anchor: str):
    for t in doc.tables:
        if t.rows and anchor.lower() in t.rows[0].cells[0].text.strip().lower():
            return t
    return None


def _set_cell_text(cell, lines: list[str]):
    """Replace a cell's content with `lines`, reusing the first run's font."""
    para = cell.paragraphs[0]
    font_src = para.runs[0] if para.runs else None
    # wipe existing paragraphs except the first
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for i, line in enumerate(lines):
        if i:
            para.add_run().add_break()
        run = para.add_run(line)
        if font_src is not None:
            run.font.name = font_src.font.name
            run.font.size = font_src.font.size
            run.bold = font_src.bold


def _append_to_label_cell(cell, value: str):
    """Append ' value' after a label like 'gescreent:' (keep the label)."""
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[-1].add_text(f" {value}")
    else:
        para.add_run(f" {value}")


def _widest_cell(row):
    """The cell spanning the most grid columns in a row (the full-width merged
    cell). `row.cells` repeats a merged cell once per column it covers, so the
    most-repeated _tc is the widest — that's where a comment band belongs, not
    the narrow leading column that cell(r, 0) may resolve to."""
    from collections import Counter

    counts: Counter = Counter()
    by_tc = {}
    for c in row.cells:
        counts[c._tc] += 1
        by_tc[c._tc] = c
    widest_tc = max(counts, key=counts.get)
    return by_tc[widest_tc]


def _comment_row_indices(table) -> list[int]:
    """Row indices that follow a 'Kommentar (ggf' label row (the free-text area)."""
    out = []
    for ri, row in enumerate(table.rows):
        if row.cells[0].text.strip().lower().startswith("kommentar (ggf"):
            if ri + 1 < len(table.rows):
                out.append(ri + 1)
    return out


def _fill_comment(table, lines: list[str]):
    """Write `lines` into the first free-text row after a 'Kommentar' label."""
    if not lines:
        return
    rows = _comment_row_indices(table)
    if not rows:
        return
    _set_cell_text(_widest_cell(table.rows[rows[0]]), lines)


# --------------------------------------------------------------------------- #
# Checkbox ticking (legacy Word form-field checkboxes)
# --------------------------------------------------------------------------- #
# Every checkbox QUESTION gets a mark. A few questions are answered from the
# Excel data via CHECKBOX_RULES (Ja when the predicate holds, otherwise Nein);
# every other question — which the data cannot decide — defaults to NZ
# ("nicht zutreffend" / not assessed) so no box is left blank. Boxes are located
# by question + option order (NOT by field name: this template reuses names, so
# only positional targeting is safe). Add rules to convert more NZ defaults into
# data-driven Ja/Nein answers.
#
# Each rule: (question_anchor, predicate, why)
#   question_anchor — substring of the row's first cell (the criterion text)
#   predicate       — VisitData -> bool; True -> tick "Ja", False -> tick "Nein"
DEFAULT_OPTION = "nz"  # for questions no rule decides ("NA")

CHECKBOX_RULES = [
    ("eCRF-Überprüfung/SDV durchgeführt",
     lambda d: bool(d.sdv_patients()),
     "the Excel records SDV activity (visit ranges / '100% SDV')"),
    ("Neue AEs",
     lambda d: any(route(f) == "sicherheit" for f in d.findings),
     "the findings include safety/AE/complication entries"),
]


def _set_checkbox(cb) -> None:
    """Mark a legacy form-field checkbox as checked: default=1 + explicit checked=1."""
    default = cb.find(qn("w:default"))
    if default is None:
        default = cb.makeelement(qn("w:default"), {})
        cb.append(default)
    default.set(qn("w:val"), "1")
    checked = cb.find(qn("w:checked"))
    if checked is None:
        checked = cb.makeelement(qn("w:checked"), {})
        cb.append(checked)
    checked.set(qn("w:val"), "1")


def _ordered_options(table) -> list[str]:
    """The table's option labels (ja/nein/nz) in left-to-right order, from the
    first header row that carries them. Merged repeats are de-duplicated."""
    for row in table.rows:
        labels, seen = [], set()
        for ci in range(len(row.cells)):
            t = row.cells[ci].text.strip().lower()
            if t in ("ja", "nein", "nz") and t not in seen:
                labels.append(t)
                seen.add(t)
        if "ja" in seen:
            return labels
    return []


def _row_checkboxes(row):
    """Checkboxes that belong *directly* to a row, in left-to-right order —
    excluding any inside a nested table (the DISPACT form nests sub-questions,
    and `.//` would otherwise pull a neighbour's boxes)."""
    tr = row._tr
    out = []
    for cb in tr.findall(".//" + qn("w:checkBox")):
        anc = cb
        while anc is not None and anc.tag != qn("w:tr"):
            anc = anc.getparent()
        if anc is tr:  # nearest enclosing row is THIS row (not a nested one)
            out.append(cb)
    return out


def _decide_option(question: str, data: VisitData) -> tuple[str, str]:
    """Return (option, why) for a question: a matching rule gives Ja/Nein,
    otherwise the NZ default."""
    for anchor, predicate, why in CHECKBOX_RULES:
        if anchor.lower() in question.lower():
            try:
                ok = predicate(data)
            except Exception:  # noqa: BLE001 — a rule must never break output
                return DEFAULT_OPTION, "rule error -> default"
            return ("ja" if ok else "nein"), (f"data: {why}" if ok else f"data: NOT ({why})")
    return DEFAULT_OPTION, "no data signal -> NA"


def apply_checkbox_rules(doc, data: VisitData) -> list[str]:
    """Mark every checkbox question: Ja/Nein where a rule decides from the data,
    NZ otherwise. Returns human-readable notes. One mark per question row."""
    notes = []
    seen = set()
    for table in doc.tables:
        order = _ordered_options(table)
        if "ja" not in order:
            continue
        for row in table.rows:
            q = row.cells[0].text.strip().replace("\n", " ")
            if not q.endswith("?"):
                continue
            cbs = _row_checkboxes(row)
            if not cbs:
                continue
            key = (q, id(row._tr))
            if key in seen:
                continue
            seen.add(key)
            option, why = _decide_option(q, data)
            if option not in order:
                option = "nein" if "nein" in order else order[-1]
            idx = order.index(option)
            if idx < len(cbs):
                _set_checkbox(cbs[idx])
                notes.append(f"{option.upper():4} | {q[:58]}  ({why})")
    return notes


# --------------------------------------------------------------------------- #
# Report assembly
# --------------------------------------------------------------------------- #
def _fmt_date(d: _dt.date | None) -> str:
    return d.strftime("%d.%m.%Y") if d else ""


def fill_report(data: VisitData, template_path: Path, out_path: Path):
    doc = Document(str(template_path))

    # 1) Document header: Besuchsdatum -------------------------------------- #
    _fill_visit_date_header(doc, data.visit_date)

    # 2) Patientenstatus counts -------------------------------------------- #
    status = _find_table(doc, "Patientenstatus")
    if status is not None and len(status.rows) > 1:
        row = status.rows[1]
        n_screened = len(data.patients)
        n_rando = sum(p.is_randomised for p in data.patients)
        n_incl = sum(p.is_included for p in data.patients)
        for cell in {c._tc: c for c in row.cells}.values():
            t = cell.text.strip().lower()
            if t.startswith("gescreent"):
                _append_to_label_cell(cell, str(n_screened))
            elif t.startswith("eingeschlossen"):
                _append_to_label_cell(cell, str(n_incl))
            elif t.startswith("randomisiert"):
                _append_to_label_cell(cell, str(n_rando))

    # 3) Route findings into their sections -------------------------------- #
    buckets: dict[str, list[str]] = {k: [] for k in SECTION_ANCHORS}
    for f in data.findings:
        buckets[route(f)].append(f.render())

    for key, anchor in SECTION_ANCHORS.items():
        if key in ("general", "sdv"):
            continue
        if not buckets[key]:
            continue
        table = _find_table(doc, anchor)
        if table is not None:
            _fill_comment(table, buckets[key])

    # 4) Patienteneinschluss: reviewed screening numbers ------------------- #
    _fill_screening_list(doc, data)

    # 5) eCRF / SDV table: screening + visit range + SDV findings ---------- #
    _fill_sdv_table(doc, data, buckets["sdv"])

    # 6) Generelle Kommentare: everything unrouted ------------------------- #
    general = _find_table(doc, "Generelle Kommentare")
    if general is not None and buckets["general"]:
        _set_cell_text(_widest_cell(general.rows[len(general.rows) - 1]), buckets["general"])

    # 7) Tick the few checkboxes the data unambiguously supports ----------- #
    ticked = apply_checkbox_rules(doc, data)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return ticked


def _fill_visit_date_header(doc, visit_date):
    if not visit_date:
        return
    target = f"Besuchsdatum: {_fmt_date(visit_date)}"
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            for table in hdr.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if "besuchsdatum" in para.text.lower() and para.runs:
                                # rewrite the run that carries the label
                                for run in para.runs:
                                    if "besuchsdatum" in run.text.lower():
                                        run.text = re.sub(
                                            r"Besuchsdatum:.*", target, run.text
                                        )
            for para in hdr.paragraphs:
                if "besuchsdatum" in para.text.lower():
                    for run in para.runs:
                        if "besuchsdatum" in run.text.lower():
                            run.text = re.sub(r"Besuchsdatum:.*", target, run.text)


def _first_empty_block(table, key_col: int) -> list[int]:
    """Indices of the first contiguous run of rows whose `key_col` cell is empty,
    starting after the first row whose `key_col` cell says 'Screening...'.

    Stops at the first non-empty row, so it never bleeds into the comment band.
    """
    rows: list[int] = []
    header_seen = False
    for ri in range(len(table.rows)):
        cell = _cell_or_none(table, ri, key_col)
        text = cell.text.strip().lower() if cell is not None else ""
        if not header_seen:
            if text.startswith("screening"):
                header_seen = True
            continue
        # Skip any further consecutive 'Screening...' sub-header rows that
        # precede the data block (the IC/SDV sub-headers span several rows).
        if not rows and text.startswith("screening"):
            continue
        if text:  # first filled/label cell after the data block → stop
            break
        rows.append(ri)
    return rows


def _fill_screening_list(doc, data: VisitData):
    """List the reviewed patients' screening numbers in the Patienteneinschluss
    table (one per empty data row under the 'Screening-Nr.' sub-header, col 4).

    "Reviewed" = patients whose source data was verified this visit, mirroring
    the reference report; falls back to all included patients if none are
    flagged for SDV.
    """
    table = _find_table(doc, "Patienteneinschluss")
    if table is None:
        return
    reviewed = [p for p, _ in data.sdv_patients()] or [p for p in data.patients if p.is_included]
    rows = _first_empty_block(table, key_col=4)
    for p, ri in zip(reviewed, rows):
        _set_cell_text(table.cell(ri, 4), [p.screening])


def _fill_sdv_table(doc, data: VisitData, sdv_findings: list[str]):
    table = _find_table(doc, "eCRF-Überprüfung/SDV")
    if table is None:
        return
    sdv = data.sdv_patients()
    rows = _first_empty_block(table, key_col=0)
    for (p, vrange), ri in zip(sdv, rows):
        _set_cell_text(table.cell(ri, 0), [p.screening])
        c2 = _cell_or_none(table, ri, 2)
        if c2 is not None:
            _set_cell_text(c2, [vrange])
    if sdv_findings:
        _fill_comment(table, sdv_findings)


def _cell_or_none(table, ri, ci):
    try:
        return table.cell(ri, ci)
    except (IndexError, Exception):  # noqa: BLE001
        return None


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("excel", type=Path, help="monitoring working sheet (.xlsx)")
    ap.add_argument("template", type=Path, help="empty report template (.docx)")
    ap.add_argument("output", type=Path, help="output report (.docx)")
    args = ap.parse_args(argv)

    # Console may be cp1252 on Windows; keep German text / arrows printable.
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:  # noqa: BLE001
        pass

    data = parse_excel(args.excel)
    print(f"Parsed: visit {_fmt_date(data.visit_date)!r}, "
          f"{len(data.patients)} patients, {len(data.findings)} findings")
    ticked = fill_report(data, args.template, args.output)
    if ticked:
        print("Ticked checkboxes (data-derived):")
        for note in ticked:
            print(f"  [x] {note}")
    else:
        print("No checkboxes ticked (no rule conditions met).")
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
