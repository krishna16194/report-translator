"""Compare a generated DISPACT report against a human-filled sample.

    python scripts/compare_report.py GENERATED.docx SAMPLE.docx

Extracts and shows side-by-side: visit date, Patientenstatus counts, the set of
ticked checkboxes (with their row question), and the filled comment/table text —
so we can see what the script populated vs. what a human filled in.
"""
import re
import sys
import zipfile
from docx import Document


def checked_checkboxes(path):
    """Return list of (question, option_label) for every ticked legacy checkbox."""
    z = zipfile.ZipFile(path)
    xml = z.read("word/document.xml").decode("utf-8", "ignore")
    out = []
    # Each ffData carries a name + checkBox; ticked == default val=1 or <w:checked val="1">
    for m in re.finditer(r"<w:ffData>([\s\S]*?)</w:ffData>", xml):
        blk = m.group(1)
        if "<w:checkBox>" not in blk:
            continue
        cb = re.search(r"<w:checkBox>([\s\S]*?)</w:checkBox>", blk).group(1)
        ticked = '<w:checked w:val="1"' in cb or '<w:default w:val="1"' in cb
        if ticked:
            name = (re.search(r'<w:name w:val="([^"]*)"', blk) or [None, "?"])[1]
            out.append(name)
    return out


def text_blocks(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells]
            if any(cells):
                rows.append(cells)
        tables.append(rows)
    return paras, tables


def header_text(path):
    doc = Document(path)
    out = []
    for s in doc.sections:
        for hdr in (s.first_page_header, s.header):
            for p in hdr.paragraphs:
                if p.text.strip():
                    out.append(p.text.strip())
            for tb in hdr.tables:
                for r in tb.rows:
                    for c in r.cells:
                        if c.text.strip():
                            out.append(c.text.strip())
    return out


def find_table(tables, anchor):
    for rows in tables:
        if rows and anchor.lower() in rows[0][0].lower():
            return rows
    return None


def show(label, gen, sample):
    print(f"\n{'='*78}\n{label}\n{'='*78}")
    print(f"[GENERATED]\n{gen}\n\n[SAMPLE]\n{sample}")


def main():
    gen_path, sample_path = sys.argv[1], sys.argv[2]
    gp, gt = text_blocks(gen_path)
    sp, st = text_blocks(sample_path)

    show("HEADER", " | ".join(header_text(gen_path)), " | ".join(header_text(sample_path)))

    gstatus = find_table(gt, "Patientenstatus")
    sstatus = find_table(st, "Patientenstatus")
    show("PATIENTENSTATUS", gstatus, sstatus)

    gcb = checked_checkboxes(gen_path)
    scb = checked_checkboxes(sample_path)
    show("TICKED CHECKBOXES (count)", f"{len(gcb)} ticked: {gcb}", f"{len(scb)} ticked: {scb}")

    # Patienteneinschluss screening list
    show("PATIENTENEINSCHLUSS table",
         find_table(gt, "Patienteneinschluss"), find_table(st, "Patienteneinschluss"))

    # eCRF/SDV
    show("eCRF / SDV table",
         find_table(gt, "eCRF"), find_table(st, "eCRF"))

    show("GENERELLE KOMMENTARE",
         find_table(gt, "Generelle Kommentare"), find_table(st, "Generelle Kommentare"))

    print(f"\n{'='*78}\nSUMMARY\n{'='*78}")
    print(f"paragraphs: generated={len(gp)} sample={len(sp)}")
    print(f"tables:     generated={len(gt)} sample={len(st)}")
    print(f"checkboxes ticked: generated={len(gcb)} sample={len(scb)}")


if __name__ == "__main__":
    main()
